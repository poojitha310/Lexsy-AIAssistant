import os
import uuid
import aiofiles
from typing import List, Dict, Optional
from pathlib import Path
import PyPDF2
import docx
from io import BytesIO
from config import settings

class DocumentService:
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md'
        }
    
    async def save_uploaded_file(self, file_content: bytes, filename: str, content_type: str) -> Dict:
        """Save uploaded file and return file info"""
        try:
            # Validate file type
            if content_type not in self.supported_types:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {content_type}",
                    "supported_types": list(self.supported_types.values())
                }
            
            # Validate file size
            if len(file_content) > settings.MAX_FILE_SIZE:
                return {
                    "success": False,
                    "error": f"File too large. Max size: {settings.MAX_FILE_SIZE / (1024*1024):.1f}MB"
                }
            
            # Generate unique filename
            file_extension = self.supported_types[content_type]
            unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            return {
                "success": True,
                "filename": unique_filename,
                "original_filename": filename,
                "file_path": str(file_path),
                "file_type": file_extension,
                "file_size": len(file_content)
            }
            
        except Exception as e:
            print(f"❌ Error saving file: {e}")
            return {"success": False, "error": str(e)}
    
    def extract_text_from_file(self, file_path: str, file_type: str) -> Dict:
        """Extract text from uploaded file"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"success": False, "error": "File not found"}
            
            extracted_text = ""
            metadata = {}
            
            if file_type == 'pdf':
                extracted_text, metadata = self._extract_from_pdf(file_path)
            elif file_type == 'docx':
                extracted_text, metadata = self._extract_from_docx(file_path)
            elif file_type in ['txt', 'md']:
                extracted_text, metadata = self._extract_from_text(file_path)
            else:
                return {"success": False, "error": f"Unsupported file type: {file_type}"}
            
            return {
                "success": True,
                "text": extracted_text,
                "metadata": metadata,
                "word_count": len(extracted_text.split()),
                "char_count": len(extracted_text)
            }
            
        except Exception as e:
            print(f"❌ Error extracting text: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_from_pdf(self, file_path: Path) -> tuple:
        """Extract text from PDF file"""
        try:
            text = ""
            metadata = {"pages": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                # Get PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "producer": pdf_reader.metadata.get('/Producer', ''),
                        "creation_date": str(pdf_reader.metadata.get('/CreationDate', '')),
                        "modification_date": str(pdf_reader.metadata.get('/ModDate', ''))
                    })
            
            return text.strip(), metadata
            
        except Exception as e:
            print(f"❌ Error extracting PDF: {e}")
            return "", {"error": str(e)}
    
    def _extract_from_docx(self, file_path: Path) -> tuple:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            if table_text:
                text_parts.append("\n--- Tables ---\n" + "\n".join(table_text))
            
            text = "\n".join(text_parts)
            
            # Extract metadata
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            # Core properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    "title": props.title or "",
                    "author": props.author or "", 
                    "subject": props.subject or "",
                    "keywords": props.keywords or "",
                    "comments": props.comments or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else ""
                })
            
            return text, metadata
            
        except Exception as e:
            print(f"❌ Error extracting DOCX: {e}")
            return "", {"error": str(e)}
    
    def _extract_from_text(self, file_path: Path) -> tuple:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "lines": len(text.split('\n')),
                "encoding": "utf-8"
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                metadata = {
                    "lines": len(text.split('\n')),
                    "encoding": "latin-1"
                }
                return text, metadata
            except Exception as e:
                print(f"❌ Error with encoding: {e}")
                return "", {"error": f"Encoding error: {e}"}
        except Exception as e:
            print(f"❌ Error extracting text: {e}")
            return "", {"error": str(e)}
    
    def get_sample_documents(self) -> List[Dict]:
        """Get sample Lexsy documents for demo"""
        return [
            {
                "filename": "lexsy-board-approval-equity-incentive.pdf",
                "original_filename": "Lexsy, Inc - Board Approval of Equity Incentive Plan.pdf",
                "file_type": "pdf",
                "content": """BOARD CONSENT OF LEXSY, INC.

WRITTEN CONSENT IN LIEU OF MEETING

The undersigned, being all of the directors of Lexsy, Inc., a Delaware corporation (the "Company"), hereby take the following actions by written consent pursuant to Section 141(f) of the Delaware General Corporation Law:

APPROVAL OF EQUITY INCENTIVE PLAN

WHEREAS, the Company desires to adopt an Equity Incentive Plan to attract, retain, and motivate employees, directors, and consultants;

WHEREAS, the Board has reviewed the proposed 2025 Equity Incentive Plan;

NOW, THEREFORE, IT IS RESOLVED:

1. The Company's 2025 Equity Incentive Plan, substantially in the form presented to the Board, is hereby approved and adopted.

2. The Plan shall reserve 1,000,000 shares of Common Stock for issuance.

3. The officers of the Company are authorized to take all actions necessary to implement the Plan.

Dated: July 15, 2025

_________________________
Sarah Chen, Director

_________________________
Alex Rodriguez, Director

_________________________
Michael Thompson, Director"""
            },
            {
                "filename": "lexsy-advisor-agreement-template.docx", 
                "original_filename": "Lexsy, Inc. - Form of Advisor Agreement.docx",
                "file_type": "docx",
                "content": """ADVISOR AGREEMENT

This Advisor Agreement ("Agreement") is entered into as of [DATE] between Lexsy, Inc., a Delaware corporation ("Company"), and [ADVISOR NAME] ("Advisor").

1. ADVISORY SERVICES
Advisor agrees to provide strategic advice and counsel to the Company in the areas of:
• AI technology development and strategy
• Venture capital and investor introductions
• Business development and partnerships
• Market positioning and competitive analysis

2. COMPENSATION
In consideration for services, Company will grant Advisor [NUMBER] shares of Restricted Stock Awards ("RSAs") subject to the terms of the Company's 2025 Equity Incentive Plan.

3. VESTING SCHEDULE
The RSAs shall vest over 24 months, with 1/24th vesting each month on the monthly anniversary of the Grant Date, subject to Advisor's continued service.

4. CONFIDENTIALITY
Advisor acknowledges that they will receive confidential information and agrees to maintain strict confidentiality.

5. TERM
This Agreement shall remain in effect for 2 years unless terminated earlier.

IN WITNESS WHEREOF, the parties have executed this Agreement.

LEXSY, INC.

By: _______________________
Name: Alex Rodriguez  
Title: Chief Executive Officer

ADVISOR:

_______________________
[Advisor Name]"""
            },
            {
                "filename": "lexsy-equity-incentive-plan.pdf",
                "original_filename": "Lexsy, Inc. - Equity Incentive Plan (EIP).pdf", 
                "file_type": "pdf",
                "content": """LEXSY, INC.
2025 EQUITY INCENTIVE PLAN

1. PURPOSE
This Plan is intended to attract, retain, and motivate employees, directors, and consultants by providing equity-based compensation opportunities.

2. SHARES AVAILABLE
Subject to adjustment, the total number of shares available for grants under this Plan is 1,000,000 shares of Common Stock.

3. ADMINISTRATION  
The Plan shall be administered by the Board of Directors or a committee designated by the Board.

4. ELIGIBLE PARTICIPANTS
Employees, directors, and consultants of the Company and its subsidiaries are eligible to participate.

5. TYPES OF AWARDS
The Plan permits the following types of awards:
• Stock Options (Incentive and Non-Qualified)
• Restricted Stock Awards (RSAs)
• Restricted Stock Units (RSUs)
• Stock Appreciation Rights

6. VESTING
Awards may vest based on:
• Continued service over time
• Achievement of performance goals
• Other conditions determined by the Administrator

7. SHARE COUNTING
For purposes of the share limit:
• Each share subject to an Option or SAR counts as 1 share
• Each share subject to RSA/RSU counts as 1 share

Current Available Shares: 985,000 (15,000 shares previously granted)

8. TERM
This Plan shall remain in effect for 10 years from the adoption date.

Adopted: July 15, 2025
Effective: July 15, 2025"""
            }
        ]
    
    def cleanup_file(self, file_path: str) -> bool:
        """Delete uploaded file"""
        try:
            file_path = Path(file_path)
            if file_path.exists():
                file_path.unlink()
                return True
            return False
        except Exception as e:
            print(f"❌ Error deleting file: {e}")
            return False